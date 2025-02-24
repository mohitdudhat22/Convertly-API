import os
import io
import subprocess
import fitz
import zipfile
from docx import Document
from PIL import Image
from fastapi import FastAPI, UploadFile, File, Form
from fastapi.responses import FileResponse
from fastapi.responses import StreamingResponse
app = FastAPI()

# Define output folder
OUTPUT_FOLDER = "Compressed"

# Ensure output folder exists
if not os.path.exists(OUTPUT_FOLDER):
    os.makedirs(OUTPUT_FOLDER)

# Define compression settings
COMPRESSION_LEVELS = {
    "low": "/screen",   # Smallest file size, lower quality
    "medium": "/ebook", # Good balance of quality and compression
    "high": "/printer"  # Higher quality, larger file size
}


@app.get("/")
def read_root():
    return {"message": "Hello, World!"}

@app.post("/compress/")
async def compress_pdf(
    file: UploadFile = File(...),
    compression_mode: str = Form(...)
):
    print(f"Received compression mode: {compression_mode}")
    print(f"Received file: {file.filename}")

    if compression_mode not in COMPRESSION_LEVELS:
        return {"error": "Invalid compression mode! Choose from: low, medium, high"}

    input_path = file.filename
    output_path = os.path.join(OUTPUT_FOLDER, file.filename)

    with open(input_path, "wb") as buffer:
        buffer.write(await file.read())

    # subprocess.call([
    #     "gs", "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.4",
    #     f"-dPDFSETTINGS={COMPRESSION_LEVELS[compression_mode]}",
    #     "-dNOPAUSE", "-dQUIET", "-dBATCH",
    #     f"-sOutputFile={output_path}", input_path
    # ])

    subprocess.call([
        "C:\\Program Files\\gs\\gs10.04.0\\bin\\gswin64c.exe", "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.4",
        f"-dPDFSETTINGS={COMPRESSION_LEVELS[compression_mode]}",
        "-dNOPAUSE", "-dQUIET", "-dBATCH",
        f"-sOutputFile={output_path}", input_path
    ])

    return FileResponse(output_path, filename=file.filename, media_type="application/pdf")



@app.post("/protect/")
async def protect_pdf(
    file: UploadFile = File(...),
    password: str = Form(...)
):
    """
    Uploads a PDF file and encrypts it with a password using PyMuPDF.
    """
    print(f"Received file: {file.filename}")
    print(f"Setting password protection.")

    input_path = file.filename
    protected_output_path = os.path.join(OUTPUT_FOLDER, f"protected_{file.filename}")

    # Save the uploaded file
    with open(input_path, "wb") as buffer:
        buffer.write(await file.read())

    # Open the PDF and encrypt it
    doc = fitz.open(input_path)
    doc.save(protected_output_path, encryption=fitz.PDF_ENCRYPT_AES_256, user_pw=password)

    return FileResponse(protected_output_path, filename=f"protected_{file.filename}", media_type="application/pdf")



@app.post("/convert/")
async def convert_pdf_to_jpg(file: UploadFile = File(...)):
    """Convert each PDF page to JPG and return as a ZIP file."""
    input_pdf = await file.read()
    doc = fitz.open(stream=input_pdf, filetype="pdf")

    zip_stream = io.BytesIO()

    with zipfile.ZipFile(zip_stream, "w", zipfile.ZIP_DEFLATED) as zipf:
        for i, page in enumerate(doc):
            # Render PDF page as an image
            pix = page.get_pixmap()
            img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

            # Convert image to in-memory JPG
            img_bytes = io.BytesIO()
            img.save(img_bytes, format="JPEG", quality=90)
            img_bytes.seek(0)

            filename = f"page_{i+1}.jpg"
            zipf.writestr(filename, img_bytes.getvalue())

    zip_stream.seek(0)

    return StreamingResponse(
        zip_stream,
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=converted_images.zip"}
    )



@app.post("/convert/pdf-to-word/")
async def convert_pdf_to_word(file: UploadFile = File(...)):
    """Convert a PDF file to a Word document and return as a downloadable file."""
    input_pdf = await file.read()
    doc = fitz.open(stream=input_pdf, filetype="pdf")

    word_doc = Document()

    for page in doc:
        text = page.get_text("text")  # Extract text from the PDF page
        word_doc.add_paragraph(text)
        word_doc.add_paragraph("\n")  # Add space between pages

    word_stream = io.BytesIO()
    word_doc.save(word_stream)
    word_stream.seek(0)

    return StreamingResponse(
        word_stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=converted.docx"}
    )